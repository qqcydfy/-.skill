#!/usr/bin/env python3
"""
周报生成器 - 生成仿真项目周报Word文档
基于原版周报格式：字体大小、加粗、段落间距、图片题注位置
"""

import argparse
import json
import os
import glob
from pathlib import Path
from typing import List, Dict, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请安装python-docx: pip install python-docx")
    exit(1)


# 样式常量（基于原版周报分析）
TITLE_SIZE = Pt(22)        # 周报标题：22pt
PROGRESS_SIZE = Pt(18)     # 进展概述：18pt
MODULE_TITLE_SIZE = Pt(16) # 模块标题（①②③）：16pt
BODY_SIZE = Pt(15)         # 正文内容：15pt
CAPTION_SIZE = Pt(10)      # 图片说明：10pt
SUBTITLE_SIZE = Pt(15)     # 子标题：15pt

# 段落间距常量（单位：twips，1pt = 20twips）
TITLE_SPACE_BEFORE = Pt(0)
TITLE_SPACE_AFTER = Pt(12)
MODULE_TITLE_SPACE_BEFORE = Pt(12)
MODULE_TITLE_SPACE_AFTER = Pt(6)
BODY_SPACE_BEFORE = Pt(0)
BODY_SPACE_AFTER = Pt(6)
CAPTION_SPACE_BEFORE = Pt(3)
CAPTION_SPACE_AFTER = Pt(6)


def set_paragraph_spacing(paragraph, space_before=None, space_after=None, line_spacing=None):
    """设置段落间距

    Args:
        paragraph: 段落对象
        space_before: 段前间距
        space_after: 段后间距
        line_spacing: 行间距
    """
    pf = paragraph.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing


class WeeklyReportGenerator:
    """周报生成器类"""

    def __init__(self, template_path: Optional[str] = None):
        """初始化生成器

        Args:
            template_path: 可选的模板文件路径
        """
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()
            self._setup_styles()

    def _setup_styles(self):
        """设置文档样式"""
        # 设置Normal样式（正文）
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = '宋体'
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        normal_style.font.size = BODY_SIZE

        # 设置Caption样式（图片说明）
        if 'Caption' in self.doc.styles:
            caption_style = self.doc.styles['Caption']
            caption_style.font.name = 'Arial'
            caption_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            caption_style.font.size = CAPTION_SIZE
            caption_style.font.italic = False

    def _add_paragraph_with_style(self, text: str, size: Pt, bold: bool = False,
                                   alignment=None, style_name: str = 'Normal',
                                   space_before=None, space_after=None):
        """添加带样式的段落

        Args:
            text: 段落文本
            size: 字体大小
            bold: 是否加粗
            alignment: 对齐方式
            style_name: 基础样式名称
            space_before: 段前间距
            space_after: 段后间距
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = size
        run.bold = bold
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        if alignment is not None:
            para.alignment = alignment

        # 设置段落间距
        if space_before is not None or space_after is not None:
            set_paragraph_spacing(para, space_before=space_before, space_after=space_after)

        return para

    def generate(self, config: Dict, output_path: str):
        """根据配置生成周报

        Args:
            config: 周报配置字典
            output_path: 输出文件路径
        """
        # 添加标题（22pt，居中，不加粗）
        title = f"{config.get('date_range', '')}周报 {config.get('author', '')}"
        self._add_paragraph_with_style(title, TITLE_SIZE, bold=False,
                                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                        space_after=TITLE_SPACE_AFTER)

        # 添加空行
        self.doc.add_paragraph("")

        # 添加进展概述（18pt，不加粗）
        self._add_paragraph_with_style("进展：", PROGRESS_SIZE, bold=False,
                                        space_after=BODY_SPACE_AFTER)

        # 进展内容（15pt，不加粗）
        if 'summary' in config:
            self._add_paragraph_with_style(config['summary'], BODY_SIZE, bold=False,
                                            space_after=BODY_SPACE_AFTER)

        # 添加各模块内容
        for idx, section in enumerate(config.get('sections', []), 1):
            self._add_section(section, idx, config.get('image_dir'))

        # 添加总结
        if 'conclusion' in config:
            self.doc.add_paragraph("")
            # 总结标题（16pt，加粗）
            self._add_paragraph_with_style("总结", MODULE_TITLE_SIZE, bold=True,
                                            space_before=MODULE_TITLE_SPACE_BEFORE,
                                            space_after=MODULE_TITLE_SPACE_AFTER)
            # 总结内容（15pt，不加粗）
            for line in config['conclusion']:
                self._add_paragraph_with_style(line, BODY_SIZE, bold=False,
                                                space_after=BODY_SPACE_AFTER)

        # 保存文档
        self.doc.save(output_path)
        print(f"周报已生成: {output_path}")

    def _add_section(self, section: Dict, idx: int, image_dir: Optional[str] = None):
        """添加一个章节

        Args:
            section: 章节配置
            idx: 章节序号
            image_dir: 图片目录路径
        """
        # 圆圈数字序号映射
        circle_numbers = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩']
        circle_num = circle_numbers[idx - 1] if idx <= len(circle_numbers) else f"({idx})"

        # 第2个及以后的模块前添加空行分隔
        if idx > 1:
            self.doc.add_paragraph("")
            self.doc.add_paragraph("")

        # 添加章节标题（16pt，加粗）
        title = section.get('title', f'模块{idx}')
        section_title = f"{circle_num}{title}"
        self._add_paragraph_with_style(section_title, MODULE_TITLE_SIZE, bold=True,
                                        space_before=MODULE_TITLE_SPACE_BEFORE,
                                        space_after=MODULE_TITLE_SPACE_AFTER)

        # 分序号计数器
        numbered_counter = 0

        # 添加章节内容
        for content in section.get('content', []):
            if isinstance(content, str):
                # 正文内容（15pt，不加粗）
                self._add_paragraph_with_style(content, BODY_SIZE, bold=False,
                                                space_after=BODY_SPACE_AFTER)
            elif isinstance(content, dict):
                # 处理子标题（15pt，加粗）
                if 'subtitle' in content:
                    self._add_paragraph_with_style(content['subtitle'], SUBTITLE_SIZE, bold=True,
                                                    space_before=MODULE_TITLE_SPACE_BEFORE,
                                                    space_after=MODULE_TITLE_SPACE_AFTER)
                # 处理分序号（1）2）3）等）
                elif 'numbered' in content:
                    self._add_paragraph_with_style(content['numbered'], SUBTITLE_SIZE, bold=True,
                                                    space_before=MODULE_TITLE_SPACE_BEFORE,
                                                    space_after=MODULE_TITLE_SPACE_AFTER)
                # 处理图片
                elif 'image' in content:
                    self._add_image(content['image'], image_dir)

        # 添加图片
        for img_info in section.get('images', []):
            self._add_image(img_info, image_dir)

    def _add_image(self, img_info: Dict, image_dir: Optional[str] = None):
        """添加图片（先图片，后题注）

        Args:
            img_info: 图片信息字典
            image_dir: 图片目录路径
        """
        # 获取图片路径
        img_path = img_info.get('path', '')
        if image_dir and not os.path.isabs(img_path):
            img_path = os.path.join(image_dir, img_path)

        # 先插入图片
        if os.path.exists(img_path):
            try:
                # 设置图片宽度为15cm
                self.doc.add_picture(img_path, width=Cm(15))
                # 图片居中
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 设置图片段落间距
                set_paragraph_spacing(last_paragraph, space_before=BODY_SPACE_BEFORE,
                                      space_after=CAPTION_SPACE_BEFORE)
                print(f"已插入图片: {img_path}")
            except Exception as e:
                print(f"插入图片失败 {img_path}: {e}")
                self._add_paragraph_with_style(f"[图片插入失败：{img_path}]", BODY_SIZE)
        else:
            print(f"图片不存在: {img_path}")
            caption = img_info.get('caption', '')
            self._add_paragraph_with_style(f"[请在此处插入图片：{caption}]", BODY_SIZE)

        # 后添加图片题注（10pt，Arial字体，居中，黑色）
        caption = img_info.get('caption', '')
        if caption:
            caption_para = self.doc.add_paragraph()
            run = caption_para.add_run(caption)
            run.font.size = CAPTION_SIZE
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 题注居中
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                caption_para.style = self.doc.styles['Caption']
            except KeyError:
                pass

            # 设置字体颜色为黑色（直接修改XML覆盖Caption样式的蓝色）
            rPr = run._element.get_or_add_rPr()
            # 移除现有的颜色设置
            for color_elem in rPr.findall(qn('w:color')):
                rPr.remove(color_elem)
            # 添加黑色颜色
            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), '000000')
            rPr.append(color_elem)

            # 设置题注段落间距
            set_paragraph_spacing(caption_para, space_before=CAPTION_SPACE_BEFORE,
                                  space_after=CAPTION_SPACE_AFTER)

        # 添加备注（15pt，不加粗）
        if 'note' in img_info:
            self._add_paragraph_with_style(f"备注：{img_info['note']}", BODY_SIZE, bold=False,
                                            space_after=BODY_SPACE_AFTER)


def scan_image_directory(image_dir: str) -> List[str]:
    """扫描图片目录，返回图片文件列表

    Args:
        image_dir: 图片目录路径

    Returns:
        图片文件路径列表
    """
    if not os.path.exists(image_dir):
        return []

    image_extensions = ['*.png', '*.jpg', '*.jpeg', '*.bmp', '*.tiff', '*.gif']
    image_files = []

    for ext in image_extensions:
        image_files.extend(glob.glob(os.path.join(image_dir, ext)))

    # 按文件名排序
    image_files.sort()
    return image_files


def auto_match_images(sections: List[Dict], image_dir: str) -> List[Dict]:
    """自动匹配图片到章节

    Args:
        sections: 章节列表
        image_dir: 图片目录路径

    Returns:
        更新后的章节列表
    """
    image_files = scan_image_directory(image_dir)
    if not image_files:
        return sections

    img_idx = 0
    for section in sections:
        # 为每个章节自动分配图片
        if 'images' not in section:
            section['images'] = []

        # 查找内容中的图片引用
        for content in section.get('content', []):
            if isinstance(content, str) and '图' in content and img_idx < len(image_files):
                # 尝试从文本中提取图号
                import re
                match = re.search(r'图\s*(\d+)', content)
                if match:
                    fig_num = int(match.group(1))
                    if fig_num <= len(image_files):
                        section['images'].append({
                            'path': image_files[fig_num - 1],
                            'caption': content
                        })
                        img_idx += 1

    return sections


def load_config(config_path: str) -> Dict:
    """加载配置文件

    Args:
        config_path: 配置文件路径

    Returns:
        配置字典
    """
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def main():
    parser = argparse.ArgumentParser(description='生成仿真项目周报')
    parser.add_argument('--config', '-c', required=True, help='配置文件路径')
    parser.add_argument('--output', '-o', default=None, help='输出文件路径')
    parser.add_argument('--template', '-t', default=None, help='模板文件路径')
    parser.add_argument('--image-dir', '-i', default=None, help='图片目录路径（覆盖配置文件中的设置）')
    parser.add_argument('--auto-images', '-a', action='store_true', help='自动匹配图片')

    args = parser.parse_args()

    # 加载配置
    config = load_config(args.config)

    # 覆盖图片目录设置
    if args.image_dir:
        config['image_dir'] = args.image_dir

    # 自动匹配图片
    if args.auto_images and 'image_dir' in config:
        config['sections'] = auto_match_images(config['sections'], config['image_dir'])

    # 生成输出文件名
    if args.output:
        output_path = args.output
    else:
        author = config.get('author', '作者')
        date_range = config.get('date_range', '日期')
        output_path = f"周报_{author}_{date_range}.docx"

    # 生成周报
    generator = WeeklyReportGenerator(args.template)
    generator.generate(config, output_path)


if __name__ == '__main__':
    main()
